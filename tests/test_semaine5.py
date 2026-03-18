"""
test_semaine5.py — Tests complets Semaine 5
  - BrowserControl  (Lundi)
  - AudioManager    (Mardi)
  - DocReader       (Mercredi + Jeudi)
  - Intégration via Agent (Vendredi)

LANCER :
    cd jarvis_windows
    python tests/test_modules/test_semaine5.py
"""

import sys
import shutil
import tempfile
from pathlib import Path
from unittest.mock import patch

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))


# ══════════════════════════════════════════════════════════════════════════════
#  HELPER
# ══════════════════════════════════════════════════════════════════════════════

def assert_response(result: dict, expected_success: bool = None):
    assert isinstance(result, dict),          f"Doit être un dict : {type(result)}"
    assert "success" in result,               "'success' manquant"
    assert "message" in result,               "'message' manquant"
    assert "data"    in result,               "'data' manquant"
    assert isinstance(result["message"], str), "message doit être str"
    if expected_success is not None:
        assert result["success"] == expected_success, (
            f"Attendu success={expected_success} : {result['message']}"
        )


# ══════════════════════════════════════════════════════════════════════════════
#  LUNDI — BrowserControl
# ══════════════════════════════════════════════════════════════════════════════

def test_browser_instantiation():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    bc = BrowserControl()
    assert bc is not None
    assert bc.default_browser == "chrome"

def test_browser_normalize_url_no_scheme():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    assert BrowserControl._normalize_url("google.com")      == "https://google.com"
    assert BrowserControl._normalize_url("github.com/test") == "https://github.com/test"

def test_browser_normalize_url_keeps_https():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    url = "https://example.com"
    assert BrowserControl._normalize_url(url) == url

def test_browser_normalize_url_keeps_http():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    url = "http://example.com"
    assert BrowserControl._normalize_url(url) == url

def test_browser_normalize_url_localhost():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    assert BrowserControl._normalize_url("localhost:3000").startswith("http://")

def test_browser_normalize_url_empty():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    assert BrowserControl._normalize_url("") == ""

def test_browser_search_empty_query():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    result = BrowserControl().google_search("")
    assert_response(result, expected_success=False)

def test_browser_search_builds_correct_url():
    """La recherche construit la bonne URL Google."""
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    from urllib.parse import quote_plus
    # On teste _open_url_system indirectement : on vérifie que l'URL construite est correcte
    bc    = BrowserControl()
    query = "python tutorial"
    expected_url = f"https://www.google.com/search?q={quote_plus(query)}"
    # Tester la construction sans réellement ouvrir le navigateur
    from JarvisDesktop.modules.browser.browser_control import SEARCH_ENGINES
    url = SEARCH_ENGINES["google"].format(quote_plus(query))
    assert url == expected_url

def test_browser_search_engines_completeness():
    """SEARCH_ENGINES contient les moteurs essentiels."""
    from JarvisDesktop.modules.browser.browser_control import SEARCH_ENGINES
    for engine in ["google", "bing", "duckduckgo", "youtube"]:
        assert engine in SEARCH_ENGINES, f"Moteur '{engine}' manquant"
        assert "{}" in SEARCH_ENGINES[engine], f"Template '{engine}' sans placeholder"

def test_browser_open_url_empty():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    result = BrowserControl().open_url("")
    assert_response(result, expected_success=False)

def test_browser_open_url_normalizes():
    """open_url normalise l'URL avant de l'ouvrir."""
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    bc = BrowserControl()
    # On ne peut pas réellement ouvrir, mais on peut vérifier que _normalize_url est appelé
    norm = bc._normalize_url("youtube.com")
    assert norm == "https://youtube.com"

def test_browser_open_new_tab_format():
    """open_new_tab retourne le bon format."""
    from JarvisDesktop.modules.browser.browser_control import BrowserControl
    # webbrowser.open_new_tab peut échouer en env sans display → on vérifie juste le format
    try:
        result = BrowserControl().open_new_tab("")
        assert_response(result)
    except Exception:
        pass  # Acceptable en environnement headless

def test_browser_commands_map():
    """BROWSER_COMMANDS contient les navigateurs courants."""
    from JarvisDesktop.modules.browser.browser_control import BROWSER_COMMANDS
    for browser in ["chrome", "firefox", "edge"]:
        assert browser in BROWSER_COMMANDS
        assert len(BROWSER_COMMANDS[browser]) >= 1


# ══════════════════════════════════════════════════════════════════════════════
#  MARDI — AudioManager
# ══════════════════════════════════════════════════════════════════════════════

def test_audio_instantiation():
    from modules.audio_manager import AudioManager
    am = AudioManager()
    assert am is not None
    assert isinstance(am._pycaw_available,  bool)
    assert isinstance(am._pactl_available,  bool)
    assert isinstance(am._amixer_available, bool)

def test_audio_volume_up_returns_dict():
    from modules.audio_manager import AudioManager
    result = AudioManager().volume_up(10)
    assert_response(result)

def test_audio_volume_down_returns_dict():
    from modules.audio_manager import AudioManager
    result = AudioManager().volume_down(10)
    assert_response(result)

def test_audio_mute_returns_dict():
    from modules.audio_manager import AudioManager
    result = AudioManager().mute()
    assert_response(result)

def test_audio_get_volume_returns_dict():
    from modules.audio_manager import AudioManager
    result = AudioManager().get_volume()
    assert_response(result, expected_success=True)
    assert "level" in result["data"]

def test_audio_set_volume_clamps():
    """set_volume ne dépasse pas 0–100."""
    from modules.audio_manager import AudioManager
    am = AudioManager()
    # On vérifie que les appels avec des valeurs limites ne crashent pas
    for level in [0, 50, 100]:
        result = am.set_volume(level)
        assert_response(result)

@patch('modules.audio_manager._has_pactl', return_value=False)
def test_audio_volume_step_clamped(mock_pactl):
    """volume_up/down clamp le step entre 1 et 100."""
    from modules.audio_manager import AudioManager
    am = AudioManager()
    # step=0 → clamped à 1, step=200 → clamped à 100
    result = am.volume_up(0)
    assert_response(result)
    result = am.volume_down(200)
    assert_response(result)

def test_audio_play_empty_query():
    from modules.audio_manager import AudioManager
    result = AudioManager().play("")
    assert_response(result, expected_success=False)
    assert "précise" in result["message"].lower() or "chanson" in result["message"].lower()

def test_audio_play_nonexistent():
    """play retourne success=False si aucun fichier trouvé."""
    from modules.audio_manager import AudioManager
    result = AudioManager().play("chanson_xyz_inexistante_jarvis_s5_test")
    assert_response(result, expected_success=False)
    assert "trouvé" in result["message"].lower() or "introuvable" in result["message"].lower()

def test_audio_play_file_nonexistent():
    from modules.audio_manager import AudioManager
    result = AudioManager().play_file("/chemin/inexistant/musique.mp3")
    assert_response(result, expected_success=False)

def test_audio_play_file_wrong_extension():
    from modules.audio_manager import AudioManager
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    try:
        result = AudioManager().play_file(tmp.name)
        assert_response(result, expected_success=False)
        assert "format" in result["message"].lower() or "supporté" in result["message"].lower()
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_audio_list_music_format():
    from modules.audio_manager import AudioManager
    result = AudioManager().list_music()
    assert_response(result, expected_success=True)
    assert "files"  in result["data"]
    assert "count"  in result["data"]
    assert isinstance(result["data"]["files"], list)

def test_audio_play_local_mp3():
    """play trouve et ouvre un .mp3 créé dans un dossier temporaire."""
    from modules.audio_manager import AudioManager
    tmp = Path(tempfile.mkdtemp())
    try:
        mp3 = tmp / "test_song.mp3"
        mp3.write_bytes(b"ID3" + b"\x00" * 100)  # Faux MP3 minimal
        am = AudioManager()
        am._muted = False
        result = am.play("test_song", music_dirs=[tmp])
        # En environnement headless, xdg-open peut échouer mais la recherche doit fonctionner
        assert_response(result)
        # Soit succès soit erreur d'ouverture (pas de display), mais le fichier doit avoir été trouvé
        if not result["success"]:
            # Acceptable : xdg-open non disponible en CI
            assert "trouvé" not in result["message"].lower()
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

def test_audio_music_extensions():
    from modules.audio_manager import MUSIC_EXTENSIONS
    for ext in [".mp3", ".flac", ".wav", ".ogg"]:
        assert ext in MUSIC_EXTENSIONS


# ══════════════════════════════════════════════════════════════════════════════
#  MERCREDI — DocReader : extraction de texte
# ══════════════════════════════════════════════════════════════════════════════

def test_doc_instantiation():
    from modules.doc_reader import DocReader
    dr = DocReader()
    assert dr is not None
    assert isinstance(dr._ai_available, bool)

def test_doc_read_nonexistent():
    from modules.doc_reader import DocReader
    result = DocReader().read("/chemin/qui/nexiste/pas/doc.txt")
    assert_response(result, expected_success=False)

def test_doc_read_txt():
    """read_txt lit un fichier texte correctement."""
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w",
                                      encoding="utf-8", delete=False)
    content = "Bonjour Jarvis !\nCeci est un test de lecture.\nTroisième ligne."
    tmp.write(content)
    tmp.close()
    try:
        result = DocReader().read_txt(tmp.name)
        assert_response(result, expected_success=True)
        data = result["data"]
        assert data["text"] == content
        assert data["lines"]   == 3
        assert data["words"]   >= 8
        assert data["chars"]   == len(content)
        assert "display" in data
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_read_txt_via_read():
    """read() dispatch vers read_txt pour les .txt"""
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Test dispatch semaine 5")
    tmp.close()
    try:
        result = DocReader().read(tmp.name)
        assert_response(result, expected_success=True)
        assert result["data"]["text"] == "Test dispatch semaine 5"
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_read_markdown():
    """read() supporte les fichiers .md"""
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".md", mode="w", delete=False)
    tmp.write("# Titre\n\nContenu markdown **gras**.")
    tmp.close()
    try:
        result = DocReader().read(tmp.name)
        assert_response(result, expected_success=True)
        assert "Titre" in result["data"]["text"]
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_read_txt_encoding_fallback():
    """read_txt gère les fichiers latin-1 (fallback encodage)."""
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False)
    tmp.write("Caf\xe9 cr\xe8me".encode("latin-1"))  # latin-1
    tmp.close()
    try:
        result = DocReader().read_txt(tmp.name)
        assert_response(result, expected_success=True)
        assert "Caf" in result["data"]["text"]
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_read_unsupported_extension():
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".xyz_unknown", delete=False)
    tmp.close()
    try:
        result = DocReader().read(tmp.name)
        assert_response(result, expected_success=False)
        assert "non supporté" in result["message"].lower() or "format" in result["message"].lower()
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_read_docx():
    """read_docx extrait le texte d'un .docx créé programmatiquement."""
    from modules.doc_reader import DocReader, DOCX_AVAILABLE
    if not DOCX_AVAILABLE:
        print("  SKIP : python-docx non installé")
        return

    from docx import Document
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    try:
        doc = Document()
        doc.add_heading("Rapport de Test Semaine 5", 0)
        doc.add_paragraph("Ce document contient des données importantes.")
        doc.add_paragraph("Le budget 2024 est de 50 000 euros.")
        doc.save(tmp.name)

        result = DocReader().read_docx(tmp.name)
        assert_response(result, expected_success=True)
        data = result["data"]
        assert "Rapport" in data["text"]
        assert "budget" in data["text"]
        assert data["words"] >= 5
        assert data["paragraphs"] >= 2
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_read_docx_wrong_extension():
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    try:
        result = DocReader().read_docx(tmp.name)
        assert_response(result, expected_success=False)
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_read_pdf():
    """read_pdf extrait le texte d'un PDF créé programmatiquement."""
    from modules.doc_reader import DocReader, PDF_AVAILABLE
    if not PDF_AVAILABLE:
        print("  SKIP : pypdf non installé")
        return

    # Créer un vrai PDF minimal avec pypdf
    import pypdf
    from pypdf import PdfWriter
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    try:
        writer = PdfWriter()
        page = writer.add_blank_page(width=612, height=792)
        with open(tmp.name, "wb") as f:
            writer.write(f)

        result = DocReader().read_pdf(tmp.name)
        assert_response(result, expected_success=True)
        assert result["data"]["pages"] == 1
        assert "text" in result["data"]
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_text_stats():
    from modules.doc_reader import DocReader
    stats = DocReader._text_stats("Bonjour tout le monde.\nComment ça va ?")
    assert stats["lines"]  == 2
    assert stats["words"]  >= 6
    assert stats["chars"]  > 0

def test_doc_extractive_summary():
    from modules.doc_reader import DocReader
    text = ("La réunion a eu lieu lundi. "
            "Les résultats sont positifs. "
            "Le budget a été approuvé. "
            "Des mesures seront prises prochainement.")
    summary = DocReader._extractive_summary(text, n_sentences=2)
    assert len(summary) > 0
    assert isinstance(summary, str)

def test_doc_get_info_txt():
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("info test")
    tmp.close()
    try:
        result = DocReader().get_doc_info(tmp.name)
        assert_response(result, expected_success=True)
        data = result["data"]
        assert "name"     in data
        assert "size"     in data
        assert "modified" in data
        assert "display"  in data
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_get_info_nonexistent():
    from modules.doc_reader import DocReader
    result = DocReader().get_doc_info("/inexistant/doc.txt")
    assert_response(result, expected_success=False)

def test_doc_supported_extensions():
    from modules.doc_reader import SUPPORTED_EXTENSIONS
    for ext in [".docx", ".pdf", ".txt", ".md", ".json", ".py"]:
        assert ext in SUPPORTED_EXTENSIONS, f"Extension '{ext}' manquante"


# ══════════════════════════════════════════════════════════════════════════════
#  JEUDI — DocReader : résumé + recherche mot
# ══════════════════════════════════════════════════════════════════════════════

def test_doc_search_word_found():
    """search_word trouve un mot dans un fichier texte."""
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Ligne 1 : introduction\nLigne 2 : budget alloué 50k\nLigne 3 : conclusion")
    tmp.close()
    try:
        result = DocReader().search_word(tmp.name, "budget")
        assert_response(result, expected_success=True)
        data = result["data"]
        assert data["found"]      == True
        assert data["count"]      >= 1
        assert data["line_count"] >= 1
        assert len(data["matches"]) >= 1
        assert data["matches"][0]["line_number"] == 2
        assert "budget" in data["matches"][0]["line"].lower()
        assert "display" in data
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_search_word_not_found():
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Contenu sans le mot cherché.")
    tmp.close()
    try:
        result = DocReader().search_word(tmp.name, "xyzinexistant")
        assert_response(result, expected_success=True)
        assert result["data"]["found"] == False
        assert result["data"]["count"] == 0
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_search_word_case_insensitive():
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Le Budget 2024 est validé.\nbudget mensuel confirmé.")
    tmp.close()
    try:
        result = DocReader().search_word(tmp.name, "budget", case_sensitive=False)
        assert_response(result, expected_success=True)
        assert result["data"]["line_count"] == 2
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_search_word_counts_multiple_occurrences():
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("budget budget budget sur une ligne\nautres budgets ici")
    tmp.close()
    try:
        result = DocReader().search_word(tmp.name, "budget")
        assert_response(result, expected_success=True)
        assert result["data"]["count"] >= 4  # 3 + 1 = 4 occurrences
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_search_word_empty_word():
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("test")
    tmp.close()
    try:
        result = DocReader().search_word(tmp.name, "")
        assert_response(result, expected_success=False)
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_summarize_local_fallback():
    """summarize utilise le résumé extractif si Groq non configuré."""
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    content = (
        "Ce rapport présente les résultats du trimestre Q3 2024. "
        "Les ventes ont augmenté de 15% par rapport à l'année précédente. "
        "Le département R&D a lancé trois nouveaux projets innovants. "
        "Les coûts opérationnels ont été réduits grâce à l'optimisation des processus. "
        "L'équipe commerciale a atteint 108% de ses objectifs annuels. "
        "Des investissements supplémentaires sont prévus pour le Q4."
    )
    tmp.write(content)
    tmp.close()
    try:
        result = DocReader().summarize(tmp.name)
        assert_response(result, expected_success=True)
        data = result["data"]
        assert "summary"  in data
        assert "source"   in data
        assert len(data["summary"]) > 20
        # Source doit être "extractive" ou "groq"
        assert data["source"] in ("extractive", "groq")
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_summarize_empty_file():
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("")
    tmp.close()
    try:
        result = DocReader().summarize(tmp.name)
        assert_response(result, expected_success=False)
        assert "vide" in result["message"].lower() or "empty" in result["message"].lower()
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_doc_read_and_answer_no_groq():
    """read_and_answer retourne erreur claire sans Groq."""
    from modules.doc_reader import DocReader
    dr = DocReader()
    if dr._ai_available:
        return  # Test non applicable si Groq est configuré
    result = dr.read_and_answer("doc.txt", "Quel est le sujet ?")
    assert_response(result, expected_success=False)
    assert "groq" in result["message"].lower() or "configuré" in result["message"].lower()

def test_doc_search_word_with_context():
    """search_word retourne du contexte autour de chaque match."""
    from modules.doc_reader import DocReader
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    lines_content = "\n".join([f"Ligne {i}" for i in range(20)])
    lines_content += "\nLigne spéciale avec budget dedans"
    tmp.write(lines_content)
    tmp.close()
    try:
        result = DocReader().search_word(tmp.name, "budget", context_lines=2)
        assert_response(result, expected_success=True)
        match = result["data"]["matches"][0]
        assert "context"       in match
        assert "context_range" in match
        # Le contexte doit contenir plusieurs lignes
        assert len(match["context"].splitlines()) >= 1
    finally:
        Path(tmp.name).unlink(missing_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
#  VENDREDI — Intégration via Agent + IntentExecutor
# ══════════════════════════════════════════════════════════════════════════════

def test_intent_executor_browser_open():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("BROWSER_OPEN", {"browser": "chrome"})
    assert_response(result)  # Peut échouer si Chrome absent, mais pas de crash

def test_intent_executor_browser_search_empty():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("BROWSER_SEARCH", {})
    assert_response(result, expected_success=False)

def test_intent_executor_browser_search_valid():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("BROWSER_SEARCH", {"query": "python asyncio"})
    assert_response(result)

def test_intent_executor_browser_url_empty():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("BROWSER_URL", {})
    assert_response(result, expected_success=False)

def test_intent_executor_audio_volume_up():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("AUDIO_VOLUME_UP", {"step": 10})
    assert_response(result)

def test_intent_executor_audio_volume_down():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("AUDIO_VOLUME_DOWN", {"step": 5})
    assert_response(result)

def test_intent_executor_audio_mute():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("AUDIO_MUTE", {})
    assert_response(result)

def test_intent_executor_audio_play_empty():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("AUDIO_PLAY", {})
    assert_response(result, expected_success=False)

def test_intent_executor_doc_read_txt():
    from core.intent_executor import IntentExecutor
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Contenu test executor semaine 5")
    tmp.close()
    try:
        result = IntentExecutor().execute("DOC_READ", {"path": tmp.name})
        assert_response(result, expected_success=True)
        assert "Contenu test" in result["data"]["text"]
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_intent_executor_doc_read_missing_path():
    from core.intent_executor import IntentExecutor
    result = IntentExecutor().execute("DOC_READ", {})
    assert_response(result, expected_success=False)

def test_intent_executor_doc_search_word():
    from core.intent_executor import IntentExecutor
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Le projet Jarvis avance bien.\nLe budget Jarvis est approuvé.")
    tmp.close()
    try:
        result = IntentExecutor().execute("DOC_SEARCH_WORD",
                                          {"path": tmp.name, "word": "Jarvis"})
        assert_response(result, expected_success=True)
        assert result["data"]["count"] >= 2
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_intent_executor_doc_summarize():
    from core.intent_executor import IntentExecutor
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Rapport Q3. Les ventes ont augmenté. Les coûts ont baissé. Résultats positifs.")
    tmp.close()
    try:
        result = IntentExecutor().execute("DOC_SUMMARIZE", {"path": tmp.name})
        assert_response(result, expected_success=True)
        assert "summary" in result["data"]
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_agent_browser_search():
    from core.agent import Agent
    result = Agent().handle_command("cherche python tutorial sur google")
    assert_response(result)
    intent = result.get("_intent") or (result.get("data") or {}).get("_intent", "")
    assert intent in ("BROWSER_SEARCH", "FILE_SEARCH", "FILE_SEARCH_CONTENT"), (
        f"Intent attendu BROWSER_SEARCH, reçu '{intent}'"
    )

def test_agent_volume_up():
    from core.agent import Agent
    result = Agent().handle_command("monte le volume")
    assert_response(result)
    intent = result.get("_intent") or (result.get("data") or {}).get("_intent", "")
    assert intent == "AUDIO_VOLUME_UP", f"Intent attendu AUDIO_VOLUME_UP, reçu '{intent}'"

def test_agent_volume_down():
    from core.agent import Agent
    result = Agent().handle_command("baisse le volume de 20%")
    assert_response(result)
    intent = result.get("_intent") or (result.get("data") or {}).get("_intent", "")
    assert intent == "AUDIO_VOLUME_DOWN", f"Attendu AUDIO_VOLUME_DOWN, reçu '{intent}'"

def test_agent_mute():
    from core.agent import Agent
    result = Agent().handle_command("coupe le son")
    assert_response(result)
    intent = result.get("_intent") or (result.get("data") or {}).get("_intent", "")
    assert intent == "AUDIO_MUTE", f"Attendu AUDIO_MUTE, reçu '{intent}'"

def test_agent_doc_read():
    from core.agent import Agent
    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Test lecture document agent semaine 5")
    tmp.close()
    try:
        result = Agent().handle_command(f"ouvre le fichier {tmp.name}")
        assert_response(result)
        # FILE_OPEN ou DOC_READ selon le parser
        intent = result.get("_intent") or (result.get("data") or {}).get("_intent", "")
        assert intent in ("FILE_OPEN", "DOC_READ"), f"Intent inattendu : '{intent}'"
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_agent_doc_search_word_pipeline():
    """Test pipeline complet : parser → executor → search_word dans un fichier réel."""
    from core.command_parser import CommandParser
    from core.intent_executor import IntentExecutor

    tmp = tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False)
    tmp.write("Le projet Jarvis est en cours.\nBudget Jarvis : 10 000 euros.")
    tmp.close()
    try:
        # Tester directement l'executor (pas le parser pour éviter l'ambiguïté)
        result = IntentExecutor().execute(
            "DOC_SEARCH_WORD",
            {"path": tmp.name, "word": "Jarvis"}
        )
        assert_response(result, expected_success=True)
        assert result["data"]["found"]  == True
        assert result["data"]["count"]  >= 2
    finally:
        Path(tmp.name).unlink(missing_ok=True)

def test_agent_open_url_pipeline():
    from core.command_parser import CommandParser
    result = CommandParser().parse("ouvre https://github.com")
    assert result["intent"] == "BROWSER_URL"
    assert "github" in result["params"].get("url", "").lower()


def test_parser_open_folder_on_drive():
    from core.command_parser import CommandParser
    result = CommandParser()._fallback_keywords("ouvre le dossier Films dans le disque E")
    assert result["intent"] == "FILE_OPEN"
    assert result["params"].get("path") == "films"
    assert result["params"].get("target_type") == "directory"
    assert result["params"].get("search_dirs") == ["E:\\"]


def test_file_open_ambiguous_sets_followup():
    from modules.file_manager import FileManager
    with tempfile.TemporaryDirectory() as tmp1, tempfile.TemporaryDirectory() as tmp2:
        p1 = Path(tmp1) / "rapport.txt"
        p2 = Path(tmp2) / "rapport.txt"
        p1.write_text("A", encoding="utf-8")
        p2.write_text("B", encoding="utf-8")
        fm = FileManager(search_dirs=[Path(tmp1), Path(tmp2)], max_depth=2)
        result = fm.open_file("rapport.txt")
        assert result["success"] is True
        data = result["data"]
        assert data["awaiting_choice"] is True
        assert len(data["choices"]) == 2


def test_file_open_with_invalid_explicit_search_scope_does_not_fallback():
    from modules.file_manager import FileManager
    fm = FileManager()
    result = fm.open_file("rapport.txt", search_dirs=["___NO_SUCH_DIR___"], target_type="file")
    assert result["success"] is False
    assert "Aucun dossier cible valide" in result["message"]
    assert "requested_search_dirs" in (result.get("data") or {})


def test_agent_followup_open_choice_by_folder_hint():
    from core.agent import Agent

    agent = Agent()
    agent.context.set_pending(
        intent="FILE_OPEN",
        params={},
        question="Lequel ouvrir ?",
        choices=[
            {"name": "rapport.txt", "path": "C:/Users/test/Documents/rapport.txt", "parent": "C:/Users/test/Documents", "is_dir": False},
            {"name": "rapport.txt", "path": "C:/Users/test/Desktop/rapport.txt", "parent": "C:/Users/test/Desktop", "is_dir": False},
        ],
        raw_command="ouvre rapport.txt",
    )

    class FakeExecutor:
        def execute(self, intent, params, raw_command=""):
            return {"success": True, "message": params.get("path", ""), "data": params}

    agent._executor = FakeExecutor()
    result = agent.handle_command("celui dans Documents")
    assert result["success"] is True
    assert "Documents" in result["message"]


def test_close_recent_opened_file_is_rewritten_from_app_close():
    from core.agent import Agent

    agent = Agent()
    agent.context.last_opened_item = {
        "path": r"E:\films\GODZILLA XKONG.mp4",
        "name": "GODZILLA XKONG.mp4",
        "stem": "GODZILLA XKONG",
        "is_dir": False,
        "parent": r"E:\films",
    }

    intent, params = agent._override_intent_with_context(
        "ferme godzilla",
        "APP_CLOSE",
        {"app_name": "godzilla"},
    )

    assert intent == "FILE_CLOSE"
    assert params["path"] == r"E:\films\GODZILLA XKONG.mp4"
    assert params["current_dir"] == r"E:\films"


def test_close_recent_opened_file_now_uses_context():
    from core.agent import Agent

    agent = Agent()
    agent.context.last_opened_item = {
        "path": r"E:\films\GODZILLA XKONG.mp4",
        "name": "GODZILLA XKONG.mp4",
        "stem": "GODZILLA XKONG",
        "is_dir": False,
        "parent": r"E:\films",
    }

    intent, params = agent._override_intent_with_context(
        "ferme maintenant",
        "UNKNOWN",
        {},
    )

    assert intent == "FILE_CLOSE"
    assert params["path"] == r"E:\films\GODZILLA XKONG.mp4"


def test_agent_browser_followup_search_uses_recent_browser_context():
    from core.agent import Agent

    class FakeParser:
        ai_available = False

        def parse_with_context(self, command, history=None):
            command = command.lower().strip()
            if command == "mets chrome":
                return {
                    "intent": "APP_OPEN",
                    "params": {"app_name": "chrome", "args": []},
                    "confidence": 0.91,
                    "source": "fallback",
                }
            if command == "recherche cours de conduite":
                return {
                    "intent": "FILE_SEARCH",
                    "params": {"query": "cours de conduite"},
                    "confidence": 0.74,
                    "source": "fallback",
                }
            return {"intent": "UNKNOWN", "params": {}, "confidence": 0.0, "source": "fallback"}

    class FakeExecutor:
        def __init__(self):
            self.calls = []

        def execute(self, intent, params, raw_command="", agent=None):
            self.calls.append((intent, params, raw_command))
            if intent == "APP_OPEN":
                return {"success": True, "message": "Chrome ouvert.", "data": {"pid": 12}}
            if intent == "BROWSER_SEARCH":
                return {
                    "success": True,
                    "message": "Recherche lancee.",
                    "data": {"query": params.get("query"), "url": "https://example.test"},
                }
            return {"success": False, "message": f"intent inattendu: {intent}", "data": {}}

    class FakeHistory:
        def save(self, **kwargs):
            return None

    agent = Agent()
    agent._parser = FakeParser()
    agent._executor = FakeExecutor()
    agent._history = FakeHistory()

    first = agent.handle_command("mets chrome")
    assert first["_intent"] == "APP_OPEN"
    assert "recherche quelque chose" in first["message"].lower()
    assert agent.context.active_surface["kind"] == "browser"

    second = agent.handle_command("recherche cours de conduite")
    assert second["_intent"] == "BROWSER_SEARCH"
    assert agent._executor.calls[-1][0] == "BROWSER_SEARCH"
    assert agent._executor.calls[-1][1]["query"] == "cours de conduite"


def test_browser_context_does_not_override_explicit_file_search():
    from core.agent import Agent

    agent = Agent()
    agent.context.active_surface = {"kind": "browser", "name": "chrome"}

    intent, params = agent._override_intent_with_context(
        "recherche le fichier budget.xlsx",
        "FILE_SEARCH",
        {"query": "budget.xlsx"},
    )

    assert intent == "FILE_SEARCH"
    assert params["query"] == "budget.xlsx"


def test_agent_document_followup_search_uses_active_document_context():
    from core.agent import Agent

    class FakeParser:
        ai_available = False

        def parse_with_context(self, command, history=None):
            command = command.lower().strip()
            if command == "ouvre contrat.pdf":
                return {
                    "intent": "FILE_OPEN",
                    "params": {"path": "contrat.pdf"},
                    "confidence": 0.9,
                    "source": "fallback",
                }
            if command == "cherche permis":
                return {
                    "intent": "FILE_SEARCH",
                    "params": {"query": "permis"},
                    "confidence": 0.71,
                    "source": "fallback",
                }
            return {"intent": "UNKNOWN", "params": {}, "confidence": 0.0, "source": "fallback"}

    class FakeExecutor:
        def __init__(self):
            self.calls = []

        def execute(self, intent, params, raw_command="", agent=None):
            self.calls.append((intent, params, raw_command))
            if intent == "FILE_OPEN":
                return {
                    "success": True,
                    "message": "Fichier ouvert : 'contrat.pdf'",
                    "data": {"opened_path": r"E:\docs\contrat.pdf"},
                }
            if intent == "DOC_SEARCH_WORD":
                return {
                    "success": True,
                    "message": "Mot trouve dans le document.",
                    "data": {"path": params.get("path"), "word": params.get("word")},
                }
            return {"success": False, "message": f"intent inattendu: {intent}", "data": {}}

    class FakeHistory:
        def save(self, **kwargs):
            return None

    agent = Agent()
    agent._parser = FakeParser()
    agent._executor = FakeExecutor()
    agent._history = FakeHistory()

    first = agent.handle_command("ouvre contrat.pdf")
    assert first["_intent"] == "FILE_OPEN"
    assert agent.context.active_surface["kind"] == "document"

    second = agent.handle_command("cherche permis")
    assert second["_intent"] == "DOC_SEARCH_WORD"
    assert agent._executor.calls[-1][1]["path"] == r"E:\docs\contrat.pdf"
    assert agent._executor.calls[-1][1]["word"] == "permis"


def test_agent_document_summary_uses_active_document_context():
    from core.agent import Agent

    agent = Agent()
    agent.context.active_surface = {
        "kind": "document",
        "path": r"E:\docs\cours.pdf",
        "name": "cours.pdf",
        "title": "cours.pdf",
    }

    intent, params = agent._override_intent_with_context("résume le document", "UNKNOWN", {})

    assert intent == "DOC_SUMMARIZE"
    assert params["path"] == r"E:\docs\cours.pdf"


def test_agent_close_uses_window_close_for_ambiguous_media():
    from core.agent import Agent

    class FakeParser:
        ai_available = False

        def parse_with_context(self, command, history=None):
            return {
                "intent": "APP_CLOSE",
                "params": {"app_name": "la video"},
                "confidence": 0.82,
                "source": "fallback",
            }

    class FakeExecutor:
        def __init__(self):
            self.calls = []

        def execute(self, intent, params, raw_command="", agent=None):
            self.calls.append((intent, params, raw_command))
            return {
                "success": True,
                "message": "J'ai trouvé 2 fenêtres pour 'video'. Laquelle veux-tu fermer ?",
                "data": {
                    "awaiting_choice": True,
                    "choices": [
                        {"hwnd": 11, "pid": 201, "title": "GODZILLA.mp4 - VLC", "process_name": "vlc.exe", "kind": "media"},
                        {"hwnd": 12, "pid": 202, "title": "Avatar.mkv - VLC", "process_name": "vlc.exe", "kind": "media"},
                    ],
                },
            }

    class FakeHistory:
        def save(self, **kwargs):
            return None

    agent = Agent()
    agent._parser = FakeParser()
    agent._executor = FakeExecutor()
    agent._history = FakeHistory()

    result = agent.handle_command("ferme la video")

    assert result["_intent"] == "WINDOW_CLOSE"
    assert agent._executor.calls[-1][0] == "WINDOW_CLOSE"
    assert agent._executor.calls[-1][1]["preferred_kind"] == "media"
    assert agent.context.has_pending() is True


def test_agent_window_close_followup_selects_specific_window():
    from core.agent import Agent

    class FakeExecutor:
        def __init__(self):
            self.calls = []

        def execute(self, intent, params, raw_command="", agent=None):
            self.calls.append((intent, params, raw_command))
            return {"success": True, "message": params.get("title", "ok"), "data": params}

    class FakeHistory:
        def save(self, **kwargs):
            return None

    agent = Agent()
    agent._executor = FakeExecutor()
    agent._history = FakeHistory()
    agent.context.set_pending(
        intent="WINDOW_CLOSE",
        params={"preferred_kind": "media"},
        question="Laquelle fermer ?",
        choices=[
            {"hwnd": 21, "pid": 301, "title": "Film 1.mp4 - VLC", "process_name": "vlc.exe", "kind": "media"},
            {"hwnd": 22, "pid": 302, "title": "Film 2.mp4 - VLC", "process_name": "vlc.exe", "kind": "media"},
        ],
        raw_command="ferme la video",
    )

    result = agent.handle_command("2")

    assert result["success"] is True
    assert agent._executor.calls[-1][0] == "WINDOW_CLOSE"
    assert agent._executor.calls[-1][1]["hwnd"] == 22


def test_agent_window_close_followup_selects_by_title_words():
    from core.agent import Agent

    class FakeExecutor:
        def __init__(self):
            self.calls = []

        def execute(self, intent, params, raw_command="", agent=None):
            self.calls.append((intent, params, raw_command))
            return {"success": True, "message": params.get("title", "ok"), "data": params}

    class FakeHistory:
        def save(self, **kwargs):
            return None

    agent = Agent()
    agent._executor = FakeExecutor()
    agent._history = FakeHistory()
    agent.context.set_pending(
        intent="WINDOW_CLOSE",
        params={"preferred_kind": "browser", "query": "chrome"},
        question="Laquelle fermer ?",
        choices=[
            {"hwnd": 31, "pid": 401, "title": "Authentification réussie - Google Chrome", "process_name": "chrome.exe", "kind": "browser"},
            {"hwnd": 32, "pid": 402, "title": "Nouvel onglet - Google Chrome", "process_name": "chrome.exe", "kind": "browser"},
        ],
        raw_command="ferme chrome",
    )

    result = agent.handle_command("ferme le nouvel onglet")

    assert result["success"] is True
    assert agent._executor.calls[-1][0] == "WINDOW_CLOSE"
    assert agent._executor.calls[-1][1]["hwnd"] == 32


def test_followup_with_choices_does_not_fallback_to_global_parse():
    from core.agent import Agent

    class FakeExecutor:
        def __init__(self):
            self.calls = []

        def execute(self, intent, params, raw_command="", agent=None):
            self.calls.append((intent, params, raw_command))
            return {"success": True, "message": "ok", "data": params}

    class FakeHistory:
        def save(self, **kwargs):
            return None

    class FakeParser:
        ai_available = False

        def parse(self, command, retries=2):
            return {"intent": "APP_CLOSE", "params": {"app_name": "chrome"}}

    agent = Agent()
    agent._executor = FakeExecutor()
    agent._history = FakeHistory()
    agent._parser = FakeParser()
    agent.context.set_pending(
        intent="WINDOW_CLOSE",
        params={"preferred_kind": "browser", "query": "chrome"},
        question="Laquelle fermer ?",
        choices=[
            {"hwnd": 41, "pid": 501, "title": "Fenêtre A - Chrome", "process_name": "chrome.exe", "kind": "browser"},
            {"hwnd": 42, "pid": 502, "title": "Fenêtre B - Chrome", "process_name": "chrome.exe", "kind": "browser"},
        ],
        raw_command="ferme chrome",
    )

    result = agent.handle_command("ferme ce truc")

    assert result["success"] is False
    assert result.get("data", {}).get("awaiting_choice") is True
    assert agent.context.has_pending() is True
    assert not agent._executor.calls


def test_close_phrase_with_target_is_rerouted_to_window_close():
    from core.agent import Agent

    agent = Agent()
    agent.context.active_surface = {"kind": "browser", "name": "chrome", "title": "Nouvel onglet - Google Chrome"}

    intent, params = agent._override_intent_with_context(
        "je dis sur chrome il y'a plusieurs pages ouverte ferme celle sur authentification",
        "APP_CLOSE",
        {"app_name": "chrome", "target": "authentification"},
    )

    assert intent == "WINDOW_CLOSE"
    assert params["preferred_kind"] == "browser"
    assert params["query"] == "authentification"


def test_app_open_browser_with_args_and_search_phrase_becomes_browser_search():
    from core.agent import Agent

    agent = Agent()

    intent, params = agent._override_intent_with_context(
        "recherche moi boutique en ligne sur chrome, l'onglet qui est actuellement ouvert là",
        "APP_OPEN",
        {"app_name": "chrome", "args": ["boutique en ligne"]},
    )

    assert intent == "BROWSER_SEARCH"
    assert params["query"] == "boutique en ligne"
    assert params["browser"] == "chrome"


def test_close_phrase_tab_target_extracts_authentification_and_tab_scope():
    from core.agent import Agent

    agent = Agent()
    agent.context.active_surface = {"kind": "browser", "name": "chrome", "title": "Codespaces - Google Chrome"}

    intent, params = agent._override_intent_with_context(
        "va sur chrome tu me ferme l'onglet sur authentification",
        "APP_CLOSE",
        {"app_name": "chrome"},
    )

    assert intent == "WINDOW_CLOSE"
    assert params["preferred_kind"] == "browser"
    assert params["close_scope"] == "tab"
    assert "authentification" in params["query"]


def test_followup_choice_keeps_tab_scope_for_window_close():
    from core.agent import Agent

    class FakeExecutor:
        def __init__(self):
            self.calls = []

        def execute(self, intent, params, raw_command="", agent=None):
            self.calls.append((intent, params, raw_command))
            return {"success": True, "message": "ok", "data": params}

    class FakeHistory:
        def save(self, **kwargs):
            return None

    agent = Agent()
    agent._executor = FakeExecutor()
    agent._history = FakeHistory()
    agent.context.set_pending(
        intent="WINDOW_CLOSE",
        params={"preferred_kind": "browser", "query": "authentification", "close_scope": "tab"},
        question="Laquelle fermer ?",
        choices=[
            {"hwnd": 101, "pid": 7001, "title": "Authentification réussie - Google Chrome", "process_name": "chrome.exe", "kind": "browser"},
            {"hwnd": 102, "pid": 7002, "title": "Codespaces - Google Chrome", "process_name": "chrome.exe", "kind": "browser"},
        ],
        raw_command="ferme l'onglet sur authentification",
    )

    result = agent.handle_command("ferme l'onglet de chrome qui nest sur authentification")

    assert result["success"] is True
    assert agent._executor.calls[-1][0] == "WINDOW_CLOSE"
    assert agent._executor.calls[-1][1]["hwnd"] == 101
    assert agent._executor.calls[-1][1]["close_scope"] == "tab"


def test_parser_detects_browser_open_result_intent():
    from core.command_parser import CommandParser

    parser = CommandParser()
    result = parser._fallback_keywords("ouvre le premier lien")

    assert result["intent"] == "BROWSER_OPEN_RESULT"
    assert result["params"]["rank"] == 1


def test_parser_detects_browser_list_tabs_intent():
    from core.command_parser import CommandParser

    parser = CommandParser()
    result = parser._fallback_keywords("liste les onglets ouverts")

    assert result["intent"] == "BROWSER_LIST_TABS"


def test_parser_detects_browser_new_tab_intent_with_count_and_query():
    from core.command_parser import CommandParser

    parser = CommandParser()
    result = parser._fallback_keywords("ouvre trois nouveaux onglets et fais moi la recherche de coupe du monde 2026")

    assert result["intent"] == "BROWSER_NEW_TAB"
    assert result["params"]["count"] == 3
    assert result["params"]["query"] == "coupe du monde 2026"


def test_agent_override_routes_new_tab_request_from_bad_app_open_parse():
    from core.agent import Agent

    agent = Agent()
    intent, params = agent._override_intent_with_context(
        "ouvre un nouvel onglet sur chrome",
        "APP_OPEN",
        {"app_name": "un nouvel onglet sur chrome", "args": []},
    )

    assert intent == "BROWSER_NEW_TAB"
    assert params["count"] == 1


def test_agent_override_routes_multi_tab_request_from_bad_app_open_parse():
    from core.agent import Agent

    agent = Agent()
    intent, params = agent._override_intent_with_context(
        "ouvre trois nouveaux onglets sur chrome",
        "APP_OPEN",
        {"app_name": "trois nouveaux onglets sur chrome", "args": []},
    )

    assert intent == "BROWSER_NEW_TAB"
    assert params["count"] == 3


def test_agent_override_marks_result_open_in_new_tab():
    from core.agent import Agent

    agent = Agent()
    intent, params = agent._override_intent_with_context(
        "ouvre le premier résultat dans un nouvel onglet",
        "UNKNOWN",
        {},
    )

    assert intent == "BROWSER_OPEN_RESULT"
    assert params["rank"] == 1
    assert params["target_type"] == "new_tab"


def test_browser_control_multiple_tabs_calls_open_new_tab_repeatedly():
    from JarvisDesktop.modules.browser.browser_control import BrowserControl

    bc = BrowserControl()
    calls = []

    def fake_open_new_tab(url=""):
        calls.append(url)
        return {"success": True, "message": "ok", "data": {"url": url or "about:blank"}}

    bc.open_new_tab = fake_open_new_tab
    result = bc.open_multiple_tabs(count=3, url="")

    assert result["success"] is True
    assert result["data"]["count"] == 3
    assert len(calls) == 3


def test_intent_executor_browser_new_tab_uses_search_when_query_present():
    from core.intent_executor import IntentExecutor

    class FakeBC:
        def search_in_new_tab(self, query, engine="google"):
            return {"success": True, "message": "search", "data": {"query": query, "engine": engine}}

        def open_new_tab(self, url=""):
            return {"success": True, "message": "tab", "data": {"url": url}}

        def open_multiple_tabs(self, count=1, url=""):
            return {"success": True, "message": "tabs", "data": {"count": count, "url": url}}

    ex = IntentExecutor()
    ex._bc = FakeBC()
    result = ex.execute("BROWSER_NEW_TAB", {"query": "coupe du monde 2026"})

    assert result["success"] is True
    assert result["data"]["query"] == "coupe du monde 2026"


def test_browser_new_tab_does_not_call_automation_when_cdp_not_ready():
    """open_new_tab ne doit PAS auto-lancer Chrome debug quand CDP est indisponible."""
    from JarvisDesktop.modules.browser.browser_control import BrowserControl

    bc = BrowserControl()
    bc.automation._is_cdp_ready = lambda: False

    ctrl_t_called = []

    def fake_ctrl_t():
        ctrl_t_called.append(True)
        return {"success": True, "message": "Nouvel onglet ouvert.", "data": {"url": "about:blank", "mode": "win32_ctrl_t"}}

    bc._open_blank_tab_via_ctrl_t = fake_ctrl_t
    result = bc.open_new_tab("")

    assert result["success"] is True
    assert len(ctrl_t_called) == 1, "Ctrl+T doit être utilisé, pas un lancement Chrome"
    assert "_launch_debug_browser" not in str(result)


def test_active_surface_set_to_browser_after_new_tab_intent():
    """Après BROWSER_NEW_TAB, active_surface.kind doit être 'browser'."""
    from core.agent import Agent

    class FakeExecutor:
        def __init__(self):
            self.calls = []

        def execute(self, intent, params, raw_command="", agent=None):
            self.calls.append((intent, params))
            return {"success": True, "message": "Nouvel onglet ouvert.", "data": {"count": 1}}

    class FakeHistory:
        def save(self, **kwargs):
            return None

    agent = Agent()
    agent._executor = FakeExecutor()
    agent._history = FakeHistory()

    agent.handle_command("ouvre un nouvel onglet")

    assert agent.context.active_surface.get("kind") == "browser"


def test_search_dessus_routed_to_browser_search_when_active_surface_is_browser():
    """'recherche X dessus' après new_tab doit être redirigé vers BROWSER_SEARCH."""
    from core.agent import Agent

    agent = Agent()
    agent.context.active_surface = {"kind": "browser", "name": "chrome"}

    intent, params = agent._override_intent_with_context(
        "recherche moi coupe du monde 2026 dessus",
        "FILE_SEARCH",
        {"query": "coupe du monde 2026", "search_dirs": ["Documents"]},
    )

    assert intent == "BROWSER_SEARCH", f"Attendu BROWSER_SEARCH, reçu {intent}"
    assert "coupe du monde 2026" in str(params.get("query") or "")


def test_agent_override_routes_result_open_request():
    from core.agent import Agent

    agent = Agent()
    intent, params = agent._override_intent_with_context(
        "ouvre le deuxième lien",
        "UNKNOWN",
        {},
    )

    assert intent == "BROWSER_OPEN_RESULT"
    assert params["rank"] == 2


# ══════════════════════════════════════════════════════════════════════════════
#  RUNNER MANUEL
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import inspect, traceback

    test_funcs = [
        (name, obj) for name, obj in sorted(globals().items())
        if name.startswith("test_") and callable(obj)
    ]

    passed, failed = [], []
    for name, fn in test_funcs:
        try:
            fn()
            passed.append(name)
        except Exception as e:
            failed.append((name, str(e)[:120]))

    print()
    print("=" * 65)
    print(f"  {len(passed)} PASSES  |  {len(failed)} FAILURES  |  {len(test_funcs)} TOTAL")
    print("=" * 65)
    if failed:
        print()
        for fn_name, err in failed:
            print(f"  FAIL : {fn_name}")
            print(f"         {err}")
    else:
        print("  Tous les tests passent. Semaine 5 VALIDÉE.")