"""
doc_reader.py — Lecture intelligente de documents
Lire Word (.docx), PDF, texte brut, RTF.
Résumé via Azure OpenAI. Recherche de mots/phrases.

SEMAINE 5 — MERCREDI + JEUDI — IMPLÉMENTATION COMPLÈTE
  Mercredi : read_docx, read_pdf, read_txt, open_doc, extract_text
  Jeudi    : summarize_document (Azure OpenAI), search_word, get_doc_info
"""

import os
import re
from pathlib import Path
from config.logger import get_logger
from config.settings import (
    GROQ_API_KEY,
    GROQ_MODEL_NAME,
    #AZURE_OPENAI_ENDPOINT,
    #AZURE_OPENAI_API_KEY,
    #AZURE_OPENAI_DEPLOYMENT_NAME,
    #AZURE_OPENAI_API_VERSION,
)

logger = get_logger(__name__)

# ── Librairies optionnelles ───────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
    logger.info("python-docx disponible.")
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx non installé. pip install python-docx")

try:
    import pypdf
    PDF_AVAILABLE = True
    logger.info("pypdf disponible.")
except ImportError:
    try:
        import PyPDF2 as pypdf
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False
        logger.warning("pypdf non installé. pip install pypdf")

# Extensions supportées par ce module
SUPPORTED_EXTENSIONS = {
    ".docx": "Word",
    ".doc":  "Word (ancien format)",
    ".pdf":  "PDF",
    ".txt":  "Texte brut",
    ".md":   "Markdown",
    ".rtf":  "RTF",
    ".csv":  "CSV",
    ".log":  "Journal",
    ".py":   "Python",
    ".json": "JSON",
    ".xml":  "XML",
    ".html": "HTML",
    ".htm":  "HTML",
    ".ini":  "Configuration",
    ".cfg":  "Configuration",
}

# Nombre de caractères maximum à extraire (pour ne pas saturer Azure OpenAI)
MAX_CHARS_SUMMARY = 8000
MAX_CHARS_DISPLAY = 3000


class DocReader:
    """
    Lecteur de documents multi-format avec résumé IA via Groq.
    Toutes les méthodes retournent { "success": bool, "message": str, "data": dict | None }
    """

    def __init__(self):
        # Vérification de la disponibilité de Groq
        self._ai_available = bool(
            GROQ_API_KEY and not GROQ_API_KEY.startswith("VOTRE") # Adapté au préfixe Groq
        )
        self._client = None
        if self._ai_available:
            self._init_ai()

    def _init_ai(self):
        try:
            from groq import Groq
            self._client = Groq(api_key=GROQ_API_KEY)
            logger.info(f"DocReader : Groq connecté ({GROQ_MODEL_NAME}) pour les résumés.")
        except ImportError:
            self._ai_available = False
            logger.warning("DocReader : SDK groq non installé.")
        except Exception as e:
            self._ai_available = False
            logger.error(f"DocReader : erreur connexion Groq : {e}")

    # ══════════════════════════════════════════════════════════════════════════
    #  MERCREDI — Extraction de texte
    # ══════════════════════════════════════════════════════════════════════════

    def read(self, path: str) -> dict:
        """
        Point d'entrée principal : détecte le format et lit le document.

        Args:
            path : chemin du document (absolu ou nom à rechercher)

        Returns:
            data contient : text (str), pages (int), words (int), chars (int), display (str)
        """
        path_obj = Path(path)

        # Si chemin non absolu → chercher dans les dossiers par défaut
        if not path_obj.is_absolute() or not path_obj.exists():
            from modules.file_manager import FileManager
            result = FileManager().search_file(path)
            if not result["success"]:
                return self._err(f"Document introuvable : '{path}'")
            files = result["data"]["files"]
            supported = [f for f in files if Path(f["path"]).suffix.lower() in SUPPORTED_EXTENSIONS]
            if not supported:
                return self._err(f"Aucun document lisible trouvé pour '{path}'.")
            if len(supported) > 1:
                return self._ok(
                    f"{len(supported)} documents correspondent à '{path}'. Lequel lire ?",
                    {"ambiguous": True, "files": supported[:5]}
                )
            path_obj = Path(supported[0]["path"])

        ext = path_obj.suffix.lower()
        logger.info(f"Lecture document : '{path_obj.name}' (type={ext})")

        if ext == ".docx":
            return self.read_docx(str(path_obj))
        elif ext == ".pdf":
            return self.read_pdf(str(path_obj))
        elif ext in (".txt", ".md", ".log", ".py", ".json", ".xml",
                     ".html", ".htm", ".csv", ".ini", ".cfg"):
            return self.read_txt(str(path_obj))
        elif ext == ".doc":
            return self._err(
                f"Le format .doc (ancien Word) n'est pas directement supporté. "
                f"Convertis '{path_obj.name}' en .docx avec Word."
            )
        else:
            return self._err(
                f"Format '{ext}' non supporté. "
                f"Formats acceptés : {', '.join(SUPPORTED_EXTENSIONS.keys())}"
            )

    def read_docx(self, path: str) -> dict:
        """
        Extrait le texte d'un fichier Word (.docx).
        Inclut paragraphes, tableaux, en-têtes et notes de bas de page.

        Args:
            path : chemin du fichier .docx
        """
        if not DOCX_AVAILABLE:
            return self._err(
                "python-docx non installé. Lance : pip install python-docx"
            )

        path_obj = Path(path)
        if not path_obj.exists():
            return self._err(f"Fichier introuvable : '{path}'")
        if path_obj.suffix.lower() != ".docx":
            return self._err(f"Ce fichier n'est pas un .docx : '{path_obj.name}'")

        logger.info(f"Lecture DOCX : '{path_obj.name}'")
        try:
            doc = DocxDocument(str(path_obj))

            # ── Extraire les paragraphes ──────────────────────────────────────
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # ── Extraire les tableaux ─────────────────────────────────────────
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_texts.append(" | ".join(cells))

            # ── Combiner ─────────────────────────────────────────────────────
            all_text = "\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n[TABLEAUX]\n" + "\n".join(table_texts)

            # ── Propriétés du document ────────────────────────────────────────
            core_props = doc.core_properties
            stats = self._text_stats(all_text)

            display_text = all_text[:MAX_CHARS_DISPLAY]
            if len(all_text) > MAX_CHARS_DISPLAY:
                display_text += f"\n\n[... {len(all_text) - MAX_CHARS_DISPLAY} caractères supplémentaires ...]"

            logger.info(
                f"DOCX lu : {stats['paragraphs']} paragraphes, "
                f"{stats['words']} mots, {stats['chars']} caractères"
            )
            return self._ok(
                f"Document '{path_obj.name}' lu : "
                f"{stats['paragraphs']} paragraphes, {stats['words']} mots.",
                {
                    "text":       all_text,
                    "display":    display_text,
                    "path":       str(path_obj),
                    "name":       path_obj.name,
                    "format":     "docx",
                    "paragraphs": stats["paragraphs"],
                    "words":      stats["words"],
                    "chars":      stats["chars"],
                    "author":     getattr(core_props, "author", "") or "",
                    "title":      getattr(core_props, "title", "")  or "",
                    "tables":     len(doc.tables),
                }
            )
        except Exception as e:
            logger.error(f"Erreur lecture DOCX : {e}")
            return self._err(f"Impossible de lire '{path_obj.name}' : {str(e)}")

    def read_pdf(self, path: str) -> dict:
        """
        Extrait le texte d'un fichier PDF, page par page.

        Args:
            path : chemin du fichier PDF
        """
        if not PDF_AVAILABLE:
            return self._err("pypdf non installé. Lance : pip install pypdf")

        path_obj = Path(path)
        if not path_obj.exists():
            return self._err(f"Fichier introuvable : '{path}'")
        if path_obj.suffix.lower() != ".pdf":
            return self._err(f"Ce fichier n'est pas un PDF : '{path_obj.name}'")

        logger.info(f"Lecture PDF : '{path_obj.name}'")
        try:
            pages_text = []
            n_pages    = 0

            with open(str(path_obj), "rb") as f:
                reader = pypdf.PdfReader(f)
                n_pages = len(reader.pages)

                for i, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text() or ""
                        text = text.strip()
                        if text:
                            pages_text.append(f"[Page {i + 1}]\n{text}")
                    except Exception as pe:
                        logger.warning(f"Page {i + 1} non lisible : {pe}")
                        pages_text.append(f"[Page {i + 1}] (texte non extractible)")

            all_text = "\n\n".join(pages_text)
            stats    = self._text_stats(all_text)

            display_text = all_text[:MAX_CHARS_DISPLAY]
            if len(all_text) > MAX_CHARS_DISPLAY:
                display_text += f"\n\n[... {n_pages} pages, {stats['chars']} caractères au total ...]"

            logger.info(f"PDF lu : {n_pages} pages, {stats['words']} mots")
            return self._ok(
                f"PDF '{path_obj.name}' lu : {n_pages} page(s), {stats['words']} mots.",
                {
                    "text":    all_text,
                    "display": display_text,
                    "path":    str(path_obj),
                    "name":    path_obj.name,
                    "format":  "pdf",
                    "pages":   n_pages,
                    "words":   stats["words"],
                    "chars":   stats["chars"],
                }
            )
        except Exception as e:
            logger.error(f"Erreur lecture PDF : {e}")
            return self._err(f"Impossible de lire le PDF '{path_obj.name}' : {str(e)}")

    def read_txt(self, path: str, encoding: str = "utf-8") -> dict:
        """
        Lit un fichier texte brut (.txt, .md, .py, .json, etc.)

        Args:
            path     : chemin du fichier
            encoding : encodage (défaut utf-8, fallback latin-1)
        """
        path_obj = Path(path)
        if not path_obj.exists():
            return self._err(f"Fichier introuvable : '{path}'")

        logger.info(f"Lecture TXT : '{path_obj.name}'")
        try:
            # Essayer utf-8 puis latin-1
            try:
                content = path_obj.read_text(encoding=encoding)
            except UnicodeDecodeError:
                content = path_obj.read_text(encoding="latin-1")

            stats = self._text_stats(content)

            display_text = content[:MAX_CHARS_DISPLAY]
            if len(content) > MAX_CHARS_DISPLAY:
                display_text += f"\n\n[... +{len(content) - MAX_CHARS_DISPLAY} caractères ...]"

            ext_type = SUPPORTED_EXTENSIONS.get(path_obj.suffix.lower(), "Texte")
            return self._ok(
                f"'{path_obj.name}' lu : {stats['lines']} lignes, {stats['words']} mots.",
                {
                    "text":    content,
                    "display": display_text,
                    "path":    str(path_obj),
                    "name":    path_obj.name,
                    "format":  path_obj.suffix.lower().lstrip("."),
                    "lines":   stats["lines"],
                    "words":   stats["words"],
                    "chars":   stats["chars"],
                    "type":    ext_type,
                }
            )
        except Exception as e:
            return self._err(f"Impossible de lire '{path_obj.name}' : {str(e)}")

    def get_doc_info(self, path: str) -> dict:
        """
        Retourne les métadonnées d'un document sans extraire tout le texte.

        Utile pour avoir un aperçu rapide (taille, pages, mots, auteur...)
        """
        path_obj = Path(path)
        if not path_obj.exists():
            return self._err(f"Fichier introuvable : '{path}'")

        ext = path_obj.suffix.lower()
        stat = path_obj.stat()
        info = {
            "name":     path_obj.name,
            "path":     str(path_obj),
            "format":   SUPPORTED_EXTENSIONS.get(ext, ext),
            "size":     stat.st_size,
            "size_str": self._format_size(stat.st_size),
            "modified": __import__("datetime").datetime.fromtimestamp(
                stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
        }

        if ext == ".docx" and DOCX_AVAILABLE:
            try:
                doc = DocxDocument(str(path_obj))
                cp  = doc.core_properties
                info.update({
                    "author":     getattr(cp, "author", "") or "",
                    "title":      getattr(cp, "title", "")  or "",
                    "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
                    "tables":     len(doc.tables),
                })
            except Exception:
                pass
        elif ext == ".pdf" and PDF_AVAILABLE:
            try:
                with open(str(path_obj), "rb") as f:
                    reader = pypdf.PdfReader(f)
                    info["pages"]  = len(reader.pages)
                    meta = reader.metadata or {}
                    info["title"]  = meta.get("/Title", "") or ""
                    info["author"] = meta.get("/Author", "") or ""
            except Exception:
                pass

        lines = [
            f"  Nom      : {info['name']}",
            f"  Format   : {info['format']}",
            f"  Taille   : {info['size_str']}",
            f"  Modifié  : {info['modified']}",
        ]
        if "pages"    in info: lines.append(f"  Pages    : {info['pages']}")
        if "author"   in info and info["author"]: lines.append(f"  Auteur   : {info['author']}")
        if "title"    in info and info["title"]:  lines.append(f"  Titre    : {info['title']}")
        info["display"] = "\n".join(lines)

        return self._ok(f"Infos document : '{path_obj.name}'", info)

    # ══════════════════════════════════════════════════════════════════════════
    #  JEUDI — Résumé IA + Recherche mot
    # ══════════════════════════════════════════════════════════════════════════

    def summarize(self, path: str, max_length: int = 300, 
              language: str = "français") -> dict:
        """
        Résume un document via Groq (Llama 8B).

        Args:
            path       : chemin ou nom du document
            max_length : longueur approximative du résumé en mots
            language   : langue du résumé ("français", "english"...)

        Si Groq n'est pas configuré → résumé extractif local.
        """
        logger.info(f"Résumé document : '{path}'")

        # 1. Extraire le texte
        read_result = self.read(path)
        if not read_result["success"]:
            return read_result

        text = read_result["data"]["text"]
        name = read_result["data"]["name"]

        if not text.strip():
            return self._err(f"Le document '{name}' est vide ou inexploitable.")

        # 2. Tronquer si trop long (Llama 8B a une fenêtre de 8k tokens, on reste prudent)
        # MAX_CHARS_SUMMARY devrait être autour de 15000-20000 pour le 8B
        text_for_api = text[:MAX_CHARS_SUMMARY]
        truncated = len(text) > MAX_CHARS_SUMMARY

        # 3. ── Résumé via Groq ───────────────────────────────────────────
        if self._ai_available and self._client:
            try:
                prompt = (
                    f"Résume ce document en {language} en environ {max_length} mots. "
                    f"Sois concis et précis. Identifie les points clés et la conclusion.\n\n"
                    f"DOCUMENT :\n{text_for_api}"
                )
                if truncated:
                    prompt += f"\n\n[Note: document tronqué pour l'API]"

                response = self._client.chat.completions.create(
                    model=GROQ_MODEL_NAME, # Utilise llama-3.1-8b-instant
                    messages=[
                        {"role": "system", 
                        "content": "Tu es un expert en synthèse. Résume clairement en restant fidèle au texte."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=800, # Un peu plus de marge pour le résumé
                )
            
                summary = response.choices[0].message.content.strip()
                logger.info(f"Résumé Groq généré : {len(summary)} caractères")

                return self._ok(
                    f"Résumé de '{name}' généré par Groq.",
                    {
                        "summary":   summary,
                        "display":   summary,
                        "source":    "groq",
                        "doc_chars": len(text),
                        "truncated": truncated,
                        "doc_name":  name,
                        **{k: v for k, v in read_result["data"].items() 
                        if k not in ("text", "display", "summary")}
                    }
                )
            except Exception as e:
                logger.error(f"Erreur résumé Groq : {e}")
                # Fallback automatique vers l'extractif ci-dessous en cas d'erreur API

        # 4. ── Résumé extractif local (fallback) ─────────────────────────────────
        summary = self._extractive_summary(text, n_sentences=8)
        note = "" if self._ai_available else " (Groq non configuré — résumé extractif)"

        return self._ok(
            f"Résumé de '{name}'{note}.",
            {
                "summary":   summary,
                "display":   summary,
                "source":    "extractive",
                "doc_chars": len(text),
                "doc_name":  name,
                "note":      "Configure GROQ_API_KEY pour des résumés IA de haute qualité.",
            }
        )

    def search_word(self, path: str, word: str,
                    case_sensitive: bool = False,
                    context_lines: int = 2) -> dict:
        """
        Recherche un mot ou une phrase dans un document.

        Args:
            path           : chemin ou nom du document
            word           : mot ou phrase à chercher
            case_sensitive : respecter la casse (défaut False)
            context_lines  : nombre de lignes de contexte autour de chaque occurrence

        Returns:
            data contient : matches (list), count (int), display (str)
        """
        word = word.strip()
        if not word:
            return self._err("Précise le mot ou la phrase à rechercher.")

        logger.info(f"Recherche '{word}' dans '{path}'")

        # Extraire le texte
        read_result = self.read(path)
        if not read_result["success"]:
            return read_result

        text = read_result["data"]["text"]
        name = read_result["data"]["name"]
        lines = text.splitlines()

        # Chercher les occurrences avec contexte
        search_word = word if case_sensitive else word.lower()
        matches = []

        for i, line in enumerate(lines):
            line_check = line if case_sensitive else line.lower()
            if search_word in line_check:
                # Extraire le contexte
                start = max(0, i - context_lines)
                end   = min(len(lines), i + context_lines + 1)
                context = lines[start:end]

                # Trouver toutes les positions dans la ligne
                occurrences = []
                pos = 0
                while True:
                    idx = line_check.find(search_word, pos)
                    if idx == -1:
                        break
                    occurrences.append(idx)
                    pos = idx + 1

                matches.append({
                    "line_number":  i + 1,
                    "line":         line.strip(),
                    "occurrences":  len(occurrences),
                    "context":      "\n".join(context),
                    "context_range": f"L{start + 1}–L{end}",
                })

        total_count = sum(m["occurrences"] for m in matches)

        if not matches:
            return self._ok(
                f"Le mot '{word}' n'a pas été trouvé dans '{name}'.",
                {"found": False, "count": 0, "matches": [], "word": word, "doc": name}
            )

        # Construire l'affichage
        lines_display = [
            f"'{word}' trouvé {total_count} fois dans '{name}' "
            f"({len(matches)} ligne(s) correspondante(s)) :",
            "─" * 60,
        ]
        for m in matches[:10]:   # max 10 résultats affichés
            lines_display.append(f"  L{m['line_number']:>4} : {m['line'][:80]}")
        if len(matches) > 10:
            lines_display.append(f"  ... et {len(matches) - 10} autre(s) occurrence(s)")

        logger.info(f"'{word}' trouvé {total_count} fois dans {name}")
        return self._ok(
            f"'{word}' trouvé {total_count} fois dans '{name}'.",
            {
                "found":       True,
                "count":       total_count,
                "line_count":  len(matches),
                "matches":     matches[:20],   # max 20 renvoyés
                "word":        word,
                "doc":         name,
                "display":     "\n".join(lines_display),
            }
        )

    def read_and_answer(self, path: str, question: str) -> dict:
        """
        Lit un document et répond à une question dessus via Groq.

        Args:
            path     : chemin du document
            question : question à poser sur le contenu
        """
        # 1. Vérification de la disponibilité de l'IA
        if not self._ai_available or not self._client:
            return self._err(
                "IA non configurée. "
                "Configure GROQ_API_KEY dans le fichier .env pour utiliser cette fonctionnalité."
            )

        logger.info(f"Q&A sur '{path}' : '{question}'")

        # 2. Lecture du document
        read_result = self.read(path)
        if not read_result["success"]:
            return read_result

        # On récupère le texte et on le tronque pour respecter la fenêtre de contexte de Llama 8B
        text = read_result["data"]["text"][:MAX_CHARS_SUMMARY]
        name = read_result["data"]["name"]

        # 3. Appel à Groq
        try:
            response = self._client.chat.completions.create(
                model=GROQ_MODEL_NAME, # Utilisation de llama-3.1-8b-instant
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Tu es un assistant expert en analyse de documents. "
                            "Réponds en te basant UNIQUEMENT sur le contenu du document fourni. "
                            "Si l'information n'est pas dans le document, dis-le clairement."
                        )
                    },
                    {
                        "role": "user",
                        "content": f"DOCUMENT : {name}\n\n{text}\n\nQUESTION : {question}"
                    }
                ],
                temperature=0.2, # Température basse pour plus de précision factuelle
                max_tokens=500,
            )
            
            answer = response.choices[0].message.content.strip()
            
            return self._ok(
                f"Réponse à '{question}' basée sur '{name}'.",
                {
                    "answer": answer, 
                    "display": answer, 
                    "question": question, 
                    "doc": name,
                    "source": "groq"
                }
            )
        except Exception as e:
            logger.error(f"Erreur Q&A Groq : {e}")
            return self._err(f"Erreur Groq lors de l'analyse : {str(e)}")
    
    # ══════════════════════════════════════════════════════════════════════════
    #  UTILITAIRES PRIVÉS
    # ══════════════════════════════════════════════════════════════════════════

    @staticmethod
    def _text_stats(text: str) -> dict:
        """Calcule les statistiques d'un texte."""
        lines = text.splitlines()
        words = len(re.findall(r'\b\w+\b', text))
        return {
            "lines":      len(lines),
            "paragraphs": len([l for l in lines if l.strip()]),
            "words":      words,
            "chars":      len(text),
        }

    @staticmethod
    def _extractive_summary(text: str, n_sentences: int = 8) -> str:
        """
        Résumé extractif simple : prend les N premières phrases non vides.
        Fallback quand Azure OpenAI n'est pas disponible.
        """
        # Découper en phrases
        sentences = re.split(r'(?<=[.!?])\s+', text.strip())
        # Filtrer les phrases trop courtes (< 30 chars) ou trop longues
        filtered = [s.strip() for s in sentences
                    if 30 <= len(s.strip()) <= 500]
        selected = filtered[:n_sentences]
        if not selected:
            # Fallback : premières lignes
            selected = [l.strip() for l in text.splitlines()
                        if len(l.strip()) > 20][:n_sentences]
        return " ".join(selected)

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 ** 2:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 ** 3:
            return f"{size_bytes / 1024 ** 2:.1f} MB"
        return f"{size_bytes / 1024 ** 3:.2f} GB"

    @staticmethod
    def _ok(message: str, data=None) -> dict:
        return {"success": True,  "message": message, "data": data}

    @staticmethod
    def _err(message: str, data=None) -> dict:
        return {"success": False, "message": message, "data": data}