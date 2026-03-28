"""
modules/word_manager.py — Contrôle profond Word
================================================
Semaine 10 — Création/édition de documents Word via python-docx.
Génération de CV, rapports, lettres. Export PDF via pywin32 (Word COM).

Dépendances :
  pip install python-docx      # Manipulation .docx
  pip install pywin32           # Export PDF via Word COM (Windows uniquement)

Fonctionnalités :
  - Créer un document Word structuré (titres, paragraphes, tableaux, images)
  - Générer un CV professionnel complet
  - Générer un rapport formaté avec sections
  - Modifier un document existant (ajouter/remplacer texte)
  - Exporter .docx → PDF via Word COM
  - Appliquer des styles cohérents (polices, couleurs, marges)
"""

from __future__ import annotations

import os
import re
import time
from datetime import datetime
from pathlib import Path
from config.logger import get_logger

logger = get_logger(__name__)

# ── Imports optionnels ────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx non installé. pip install python-docx")

try:
    import win32com.client as win32
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False
    logger.debug("pywin32 non installé — export PDF désactivé. pip install pywin32")

# Dossier de sortie par défaut
DEFAULT_OUTPUT_DIR = Path.home() / "Documents" / "Jarvis"


class WordManager:
    """
    Gestionnaire Word complet.
    Toutes les méthodes retournent { success, message, data }.
    """

    def __init__(self, output_dir: str = None):
        self._output_dir = Path(output_dir) if output_dir else DEFAULT_OUTPUT_DIR
        self._output_dir.mkdir(parents=True, exist_ok=True)

    # ══════════════════════════════════════════════════════════════════════════
    #  CRÉATION DE DOCUMENTS
    # ══════════════════════════════════════════════════════════════════════════

    def create_document(
        self,
        title: str,
        sections: list[dict],
        filename: str = None,
        style: str = "professionnel",
        output_dir: str = None,
        open_after: bool = True,
    ) -> dict:
        """
        Crée un document Word structuré depuis zéro.

        Args:
            title    : Titre principal du document
            sections : Liste de sections, chacune avec :
                       {"heading": str, "content": str|list, "level": int (1-3),
                        "table": [[row1], [row2]], "bullet": bool}
            filename : Nom du fichier (sans extension) — auto-généré si absent
            style    : "professionnel" | "simple" | "rapport"
            output_dir : Répertoire de sortie (défaut : ~/Documents/Jarvis)
            open_after : Ouvrir le fichier après création

        Returns:
            data contient : path (str), name (str), pages_est (int)
        """
        if not DOCX_AVAILABLE:
            return self._err("python-docx non installé. pip install python-docx")

        out_dir = Path(output_dir) if output_dir else self._output_dir
        out_dir.mkdir(parents=True, exist_ok=True)

        if not filename:
            safe = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
            filename = f"{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        filepath = out_dir / f"{filename}.docx"

        try:
            doc = Document()
            self._apply_document_style(doc, style)
            self._add_title(doc, title, style)

            for section in sections:
                heading  = section.get("heading", "")
                content  = section.get("content", "")
                level    = min(max(int(section.get("level", 2)), 1), 3)
                table    = section.get("table")
                bullets  = section.get("bullets")

                if heading:
                    doc.add_heading(heading, level=level)

                if isinstance(content, list):
                    for line in content:
                        doc.add_paragraph(str(line))
                elif content:
                    doc.add_paragraph(str(content))

                if bullets:
                    for item in bullets:
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(str(item))

                if table:
                    self._add_table(doc, table)

            doc.save(str(filepath))
            logger.info(f"Document créé : {filepath.name}")

            if open_after:
                self._open_file(str(filepath))

            return self._ok(
                f"Document '{filepath.name}' créé avec succès.",
                {
                    "path":      str(filepath),
                    "name":      filepath.name,
                    "sections":  len(sections),
                    "format":    "docx",
                    "opened":    open_after,
                }
            )
        except Exception as e:
            logger.error(f"Erreur création document : {e}")
            return self._err(f"Impossible de créer le document : {e}")

    def create_cv(
        self,
        info: dict,
        filename: str = None,
        output_dir: str = None,
        open_after: bool = True,
    ) -> dict:
        """
        Génère un CV professionnel complet au format .docx.

        Args:
            info : dict avec les clés :
              - name       : str  — Prénom Nom
              - title      : str  — Titre/Poste visé
              - email      : str
              - phone      : str
              - address    : str  (optionnel)
              - linkedin   : str  (optionnel)
              - github     : str  (optionnel)
              - summary    : str  — Résumé professionnel (3-4 phrases)
              - experience : list de dict {company, role, period, description, bullets}
              - education  : list de dict {institution, degree, period, description}
              - skills     : dict {catégorie: [compétences]}
              - languages  : list de dict {language, level}
              - certifs    : list de str (optionnel)
              - projects   : list de dict {name, description, tech} (optionnel)
            filename   : nom fichier sans extension (défaut : CV_Nom_Date)
            output_dir : dossier de sortie
            open_after : ouvrir après création

        Returns:
            data : path, name, sections_count
        """
        if not DOCX_AVAILABLE:
            return self._err("python-docx non installé. pip install python-docx")

        name_str  = info.get("name", "Candidat")
        out_dir   = Path(output_dir) if output_dir else self._output_dir
        out_dir.mkdir(parents=True, exist_ok=True)

        if not filename:
            safe_name = re.sub(r'[^\w]', '_', name_str)
            filename  = f"CV_{safe_name}_{datetime.now().strftime('%Y%m%d')}"
        filepath = out_dir / f"{filename}.docx"

        try:
            doc = Document()
            self._apply_cv_style(doc)

            # ── En-tête : Nom + Contact ───────────────────────────────────────
            self._add_cv_header(doc, info)

            # ── Résumé ────────────────────────────────────────────────────────
            if info.get("summary"):
                doc.add_heading("Profil", level=1)
                p = doc.add_paragraph(info["summary"])
                p.style.font.size = Pt(10)

            # ── Expériences ───────────────────────────────────────────────────
            if info.get("experience"):
                doc.add_heading("Expérience Professionnelle", level=1)
                for exp in info["experience"]:
                    self._add_cv_experience(doc, exp)

            # ── Formation ─────────────────────────────────────────────────────
            if info.get("education"):
                doc.add_heading("Formation", level=1)
                for edu in info["education"]:
                    self._add_cv_education(doc, edu)

            # ── Compétences ───────────────────────────────────────────────────
            if info.get("skills"):
                doc.add_heading("Compétences", level=1)
                self._add_cv_skills(doc, info["skills"])

            # ── Projets ───────────────────────────────────────────────────────
            if info.get("projects"):
                doc.add_heading("Projets", level=1)
                for proj in info["projects"]:
                    p_title = doc.add_paragraph()
                    r = p_title.add_run(proj.get("name", ""))
                    r.bold = True
                    if proj.get("tech"):
                        p_title.add_run(f" — {proj['tech']}")
                    if proj.get("description"):
                        doc.add_paragraph(proj["description"])

            # ── Langues ───────────────────────────────────────────────────────
            if info.get("languages"):
                doc.add_heading("Langues", level=1)
                for lang in info["languages"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(lang.get("language", "")).bold = True
                    level_str = lang.get("level", "")
                    if level_str:
                        p.add_run(f" — {level_str}")

            # ── Certifications ────────────────────────────────────────────────
            if info.get("certifs"):
                doc.add_heading("Certifications", level=1)
                for cert in info["certifs"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(str(cert))

            doc.save(str(filepath))
            logger.info(f"CV créé : {filepath.name}")

            if open_after:
                self._open_file(str(filepath))

            return self._ok(
                f"CV de '{name_str}' créé : '{filepath.name}'.",
                {
                    "path":    str(filepath),
                    "name":    filepath.name,
                    "format":  "docx",
                    "opened":  open_after,
                }
            )
        except Exception as e:
            logger.error(f"Erreur création CV : {e}")
            return self._err(f"Impossible de créer le CV : {e}")

    def create_report(
        self,
        title: str,
        content: str | dict,
        filename: str = None,
        output_dir: str = None,
        open_after: bool = True,
    ) -> dict:
        """
        Génère un rapport professionnel structuré.

        Args:
            title   : Titre du rapport
            content : str (texte libre) ou dict avec sections structurées
                      {"introduction": str, "body": [{heading, text}], "conclusion": str}
            filename, output_dir, open_after : idem create_document
        """
        if not DOCX_AVAILABLE:
            return self._err("python-docx non installé. pip install python-docx")

        out_dir = Path(output_dir) if output_dir else self._output_dir
        out_dir.mkdir(parents=True, exist_ok=True)

        if not filename:
            safe = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
            filename = f"Rapport_{safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        filepath = out_dir / f"{filename}.docx"

        try:
            doc = Document()
            self._apply_document_style(doc, "rapport")

            # Page de titre
            doc.add_heading(title, level=0)
            p_date = doc.add_paragraph(f"Généré le {datetime.now().strftime('%d %B %Y')}")
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

            if isinstance(content, str):
                # Texte libre — découper en paragraphes
                paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                for i, para in enumerate(paragraphs):
                    if i == 0:
                        doc.add_heading("Introduction", level=1)
                    doc.add_paragraph(para)
            elif isinstance(content, dict):
                if content.get("introduction"):
                    doc.add_heading("Introduction", level=1)
                    doc.add_paragraph(content["introduction"])
                for section in (content.get("body") or []):
                    if section.get("heading"):
                        doc.add_heading(section["heading"], level=2)
                    if section.get("text"):
                        doc.add_paragraph(section["text"])
                    if section.get("bullets"):
                        for item in section["bullets"]:
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(str(item))
                    if section.get("table"):
                        self._add_table(doc, section["table"])
                if content.get("conclusion"):
                    doc.add_heading("Conclusion", level=1)
                    doc.add_paragraph(content["conclusion"])

            doc.save(str(filepath))
            logger.info(f"Rapport créé : {filepath.name}")

            if open_after:
                self._open_file(str(filepath))

            return self._ok(
                f"Rapport '{filepath.name}' créé.",
                {"path": str(filepath), "name": filepath.name, "opened": open_after}
            )
        except Exception as e:
            return self._err(f"Impossible de créer le rapport : {e}")

    # ══════════════════════════════════════════════════════════════════════════
    #  ÉDITION
    # ══════════════════════════════════════════════════════════════════════════

    def edit_document(
        self,
        path: str,
        action: str,
        content: str = "",
        search: str = "",
        replace: str = "",
        heading: str = "",
        level: int = 2,
    ) -> dict:
        """
        Modifie un document Word existant.

        Actions :
          "append"  — ajouter du contenu à la fin
          "replace" — remplacer search par replace dans tout le document
          "add_heading" — ajouter un titre + contenu
          "add_table"   — content doit être une liste de lignes "col1|col2|..."
        """
        if not DOCX_AVAILABLE:
            return self._err("python-docx non installé.")

        path_obj = Path(path)
        if not path_obj.exists():
            return self._err(f"Fichier introuvable : '{path}'")

        try:
            doc = Document(str(path_obj))

            if action == "append":
                if heading:
                    doc.add_heading(heading, level=level)
                doc.add_paragraph(content)
                msg = f"Contenu ajouté à '{path_obj.name}'."

            elif action == "replace":
                if not search:
                    return self._err("Précise le texte à remplacer (search).")
                count = 0
                for para in doc.paragraphs:
                    if search in para.text:
                        for run in para.runs:
                            if search in run.text:
                                run.text = run.text.replace(search, replace)
                                count += 1
                msg = f"{count} occurrence(s) de '{search}' remplacée(s) dans '{path_obj.name}'."

            elif action == "add_heading":
                doc.add_heading(heading or content, level=level)
                if content and heading:
                    doc.add_paragraph(content)
                msg = f"Titre '{heading}' ajouté à '{path_obj.name}'."

            elif action == "add_table":
                rows = [line.split("|") for line in content.splitlines() if line.strip()]
                if rows:
                    self._add_table(doc, rows)
                msg = f"Tableau ajouté à '{path_obj.name}'."

            else:
                return self._err(f"Action inconnue : '{action}'. Utilise : append, replace, add_heading, add_table.")

            doc.save(str(path_obj))
            return self._ok(msg, {"path": str(path_obj), "action": action})

        except Exception as e:
            return self._err(f"Erreur édition document : {e}")

    # ══════════════════════════════════════════════════════════════════════════
    #  EXPORT PDF
    # ══════════════════════════════════════════════════════════════════════════

    def export_to_pdf(self, path: str, output_path: str = None) -> dict:
        """
        Exporte un .docx en PDF via Word COM (nécessite Microsoft Word installé).
        Fallback : retourne un message si pywin32 non disponible.

        Args:
            path        : chemin du fichier .docx source
            output_path : chemin PDF de sortie (défaut : même dossier, même nom)
        """
        path_obj = Path(path)
        if not path_obj.exists():
            return self._err(f"Fichier introuvable : '{path}'")

        if not output_path:
            pdf_path = path_obj.with_suffix(".pdf")
        else:
            pdf_path = Path(output_path)
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        if not WIN32_AVAILABLE:
            # Fallback : tenter via LibreOffice si disponible
            return self._export_pdf_libreoffice(path_obj, pdf_path)

        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(path_obj.resolve()))
            doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)  # 17 = PDF
            doc.Close()
            word.Quit()
            logger.info(f"PDF exporté : {pdf_path.name}")
            return self._ok(
                f"PDF exporté : '{pdf_path.name}'.",
                {"path": str(pdf_path), "name": pdf_path.name, "source": str(path_obj)}
            )
        except Exception as e:
            logger.error(f"Erreur export PDF Word COM : {e}")
            return self._err(f"Export PDF échoué : {e}. Vérifie que Word est installé.")

    def _export_pdf_libreoffice(self, docx_path: Path, pdf_path: Path) -> dict:
        """Fallback LibreOffice pour export PDF."""
        import subprocess, shutil
        lo = shutil.which("soffice") or shutil.which("libreoffice")
        if not lo:
            return self._err(
                "Export PDF nécessite pywin32 (pip install pywin32) ou LibreOffice. "
                "Aucun des deux n'est disponible."
            )
        try:
            subprocess.run(
                [lo, "--headless", "--convert-to", "pdf",
                 "--outdir", str(pdf_path.parent), str(docx_path)],
                timeout=30, check=True,
                capture_output=True,
            )
            generated = pdf_path.parent / (docx_path.stem + ".pdf")
            if generated != pdf_path and generated.exists():
                generated.rename(pdf_path)
            return self._ok(
                f"PDF exporté via LibreOffice : '{pdf_path.name}'.",
                {"path": str(pdf_path), "name": pdf_path.name}
            )
        except Exception as e:
            return self._err(f"Export PDF LibreOffice échoué : {e}")

    # ══════════════════════════════════════════════════════════════════════════
    #  HELPERS STYLE ET MISE EN PAGE
    # ══════════════════════════════════════════════════════════════════════════

    def _apply_document_style(self, doc, style: str = "professionnel"):
        """Applique les marges et styles de base au document."""
        from docx.shared import Cm
        section = doc.sections[0]
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # Style de police par défaut
        style_normal = doc.styles["Normal"]
        style_normal.font.name = "Calibri"
        style_normal.font.size = Pt(11)

    def _apply_cv_style(self, doc):
        """Style spécifique CV — marges réduites, police professionnelle."""
        from docx.shared import Cm
        section = doc.sections[0]
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        style_normal = doc.styles["Normal"]
        style_normal.font.name = "Calibri"
        style_normal.font.size = Pt(10)

    def _add_title(self, doc, title: str, style: str = "professionnel"):
        """Ajoute le titre principal du document."""
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Bleu foncé

    def _add_cv_header(self, doc, info: dict):
        """Bloc en-tête du CV : nom, titre, contacts."""
        # Nom en grand
        h = doc.add_heading(info.get("name", ""), level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # Titre/Poste visé
        if info.get("title"):
            p = doc.add_paragraph(info["title"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        # Contacts en une ligne
        contacts = []
        if info.get("email"):   contacts.append(f"✉ {info['email']}")
        if info.get("phone"):   contacts.append(f"☎ {info['phone']}")
        if info.get("address"): contacts.append(f"📍 {info['address']}")
        if info.get("linkedin"):contacts.append(f"LinkedIn: {info['linkedin']}")
        if info.get("github"):  contacts.append(f"GitHub: {info['github']}")

        if contacts:
            p = doc.add_paragraph(" | ".join(contacts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

        # Ligne de séparation
        self._add_horizontal_rule(doc)

    def _add_cv_experience(self, doc, exp: dict):
        """Ajoute une entrée d'expérience au CV."""
        # Ligne : Rôle (gras) — Entreprise — Période
        p = doc.add_paragraph()
        r = p.add_run(exp.get("role", ""))
        r.bold = True
        r.font.size = Pt(11)

        company = exp.get("company", "")
        period  = exp.get("period", "")
        if company or period:
            suffix = f" — {company}" if company else ""
            if period:
                suffix += f" ({period})"
            p.add_run(suffix).font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        # Description
        if exp.get("description"):
            p2 = doc.add_paragraph(exp["description"])
            p2.paragraph_format.left_indent = Cm(0.5)

        # Bullets de réalisations
        for bullet in (exp.get("bullets") or []):
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.left_indent = Cm(0.5)
            pb.add_run(str(bullet))

        doc.add_paragraph()  # Espace

    def _add_cv_education(self, doc, edu: dict):
        """Ajoute une entrée de formation au CV."""
        p = doc.add_paragraph()
        r = p.add_run(edu.get("degree", ""))
        r.bold = True

        institution = edu.get("institution", "")
        period      = edu.get("period", "")
        if institution or period:
            suffix = f" — {institution}" if institution else ""
            if period:
                suffix += f" ({period})"
            p.add_run(suffix).font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        if edu.get("description"):
            p2 = doc.add_paragraph(edu["description"])
            p2.paragraph_format.left_indent = Cm(0.5)

        doc.add_paragraph()

    def _add_cv_skills(self, doc, skills: dict):
        """Ajoute un bloc de compétences organisé par catégories."""
        for category, items in skills.items():
            p = doc.add_paragraph()
            p.add_run(f"{category} : ").bold = True
            items_str = ", ".join(items) if isinstance(items, list) else str(items)
            p.add_run(items_str)

    def _add_table(self, doc, rows: list[list]):
        """Ajoute un tableau stylisé au document."""
        if not rows:
            return
        n_cols = max(len(r) for r in rows)
        table  = doc.add_table(rows=0, cols=n_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            row = table.add_row()
            for j, cell_text in enumerate(row_data):
                if j < n_cols:
                    cell = row.cells[j]
                    cell.text = str(cell_text)
                    if i == 0:  # En-tête : gras + fond bleu
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                        self._set_cell_bg(cell, "1F497D")
                        for run in cell.paragraphs[0].runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def _add_horizontal_rule(self, doc):
        """Ajoute une ligne de séparation horizontale."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _set_cell_bg(cell, hex_color: str):
        """Définit la couleur de fond d'une cellule."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    @staticmethod
    def _open_file(path: str):
        """Ouvre le fichier avec l'application par défaut."""
        try:
            os.startfile(path)
        except Exception:
            import subprocess, sys
            if sys.platform == "darwin":
                subprocess.Popen(["open", path])
            else:
                subprocess.Popen(["xdg-open", path])

    @staticmethod
    def _ok(message: str, data=None) -> dict:
        return {"success": True,  "message": message, "data": data}

    @staticmethod
    def _err(message: str, data=None) -> dict:
        return {"success": False, "message": message, "data": data}